from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from typing import List

from docx import Document

from langchain_core.embeddings import Embeddings
from langchain_core.documents import Document as LCDocument
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings

from .settings import settings


# ========================= Извлечение и очистка =========================

_RE_SOFT = re.compile(r"[\u00AD\u200B\u200C\u200D]")
_RE_MULTI_WS = re.compile(r"\s{2,}")
_RE_MANY_DASHES = re.compile(r"^\s*[-–—]{3,}\s*$")
_RE_ONLY_BORDERS = re.compile(r"^[\|\+\=\-\.]{3,}$")
_RE_WATERMARK = re.compile(r"ОГРАЖДАЮЩАЯ\s+АКТУАЛЯЦИЯ", re.IGNORECASE)

_RE_ANY_NL = re.compile(r"(?:\r\n|\r|\n|\\r\\n|\\n|\\r)+")
_RE_CTRL   = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")

_ROLE_MAP = [
    (re.compile(r"\bпоставщик\w*\b", re.I), "исполнитель"),
    (re.compile(r"\bпродавец\w*\b", re.I), "исполнитель"),
    (re.compile(r"\bподрядчик\w*\b", re.I), "исполнитель"),

    (re.compile(r"\bпокупател\w*\b", re.I), "заказчик"),
    (re.compile(r"\bклиент\w*\b", re.I), "заказчик"),
]

def _strip_weird(s: str) -> str:
    s = unicodedata.normalize("NFC", s)
    s = s.replace("\xa0", " ")
    s = _RE_SOFT.sub("", s)
    return s

def _squash_newlines_to_space(s: str) -> str:
    if not s:
        return ""
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = s.replace("\n\n", "\n").replace("\r", "\n")
    s = _RE_ANY_NL.sub(" ", s)     
    s = _RE_CTRL.sub(" ", s)      
    s = _RE_MULTI_WS.sub(" ", s)
    return s.strip()

def extract_all_text_with_fallback(path: str) -> str:
    doc = Document(path)
    joined = doc.paragraphs[0].text.strip()
    return _strip_weird(joined)

def _expand_query(q: str) -> str:
    extras = []
    low = q.lower()

    for pat, canon in _ROLE_MAP:
        if pat.search(low):
            extras.append(canon)
    if extras:
        q = f"{q} " + " ".join(sorted(set(extras)))
    return q

def clean_for_search(raw_text: str) -> str:
    t = _strip_weird(raw_text)

    if len(_RE_WATERMARK.findall(t)) >= 5:
        m = _RE_WATERMARK.search(t)
        if m:
            t = t[:m.start()]

    clean_lines: List[str] = []
    for ln in t.split("\\n\\n"):
        s = ln.strip()

        if not s:
            continue
        if _RE_MANY_DASHES.match(s) or _RE_ONLY_BORDERS.match(s):
            continue
        if _RE_ONLY_BORDERS.match(s):
            continue
        if _RE_WATERMARK.search(s):
            continue
        s = re.sub(r'(\| )*', '', s)
        s = re.sub(r'\|\\n\d+', '', s)
        s = re.sub(r'(\|---)*', '', s)
        s = re.sub(r'\|*', '', s)
        s = re.sub(r'\\n', ' ', s)
        s = re.sub(r'_+', '', s)       
        clean_lines.append(s)
    flat = "\n\n".join(clean_lines)
    return flat


# ========================= Чанкинг =========================

def _slice_with_overlap(s: str, chunk_size: int, overlap: int) -> List[str]:
    if chunk_size <= 0:
        return [s] if s else []
    res = []
    i = 0
    n = len(s)
    while i < n:
        j = min(n, i + chunk_size)
        res.append(s[i:j])
        if j == n:
            break
        i = max(j - overlap, i + 1)
    return res

def chunk_by_double_newline(
    text: str,
    chunk_size: int,
    overlap: int,
    min_chunk_chars: int = 300,
    cross_block_overlap: bool = True,
) -> List[LCDocument]:
 
    raw_blocks = [b for b in re.split(r"\n{2,}", text) if b.strip()]
    blocks = [re.sub(r"\n+", " ", b.strip()) for b in raw_blocks]
    merged: List[str] = []
    acc = ""
    for b in blocks:
        if not acc:
            acc = b
            continue
        if len(acc) < min_chunk_chars:
            acc = f"{acc} {b}".strip()
        else:
            merged.append(acc)
            acc = b
    if acc:
        merged.append(acc)

    docs = [LCDocument(page_content=m, metadata={"chunk_id": cid}) for cid, m in enumerate(merged)]
    docs: List[LCDocument] = []

    if cross_block_overlap:
        joined = "\n\n".join(merged)
        parts = _slice_with_overlap(joined, chunk_size, overlap)
        for i, p in enumerate(parts):
            docs.append(LCDocument(page_content=p.strip(), metadata={"chunk_id": i}))
        return docs

    cid = 0
    for m in merged:
        if len(m) <= max(1, chunk_size):
            docs.append(LCDocument(page_content=m, metadata={"chunk_id": cid}))
            cid += 1
        else:
            for p in _slice_with_overlap(m, chunk_size, overlap):
                docs.append(LCDocument(page_content=p, metadata={"chunk_id": cid}))
                cid += 1
    return docs


# ========================= Embeddings =========================

def get_embeddings() -> Embeddings:
    model_name = settings.EMBED_MODEL or "intfloat/multilingual-e5-small"

    return HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs={
            "device": "cpu",      
        },
        encode_kwargs={
            "normalize_embeddings": True,  
        },
    )


# ========================= Индексация =========================

def _is_noisy(txt: str) -> bool:
    if not txt:
        return True
    alpha = sum(ch.isalpha() for ch in txt)
    if alpha / max(1, len(txt)) < 0.2:
        return True
    toks = [t for t in re.split(r"\W+", txt) if t]
    if not toks:
        return True
    from collections import Counter
    _, cnt = Counter(toks).most_common(1)[0]
    if cnt >= max(6, int(len(toks) * 0.6)):
        return True
    return False

def build_index_from_docx(file_id: str, uploaded_path: str) -> str:
    base = Path(settings.STORAGE_DIR) / file_id
    base.mkdir(parents=True, exist_ok=True)
    raw = extract_all_text_with_fallback(uploaded_path)

    cleaned = clean_for_search(raw)
    (base / "document.txt").write_text(cleaned, encoding="utf-8")
    docs_all = chunk_by_double_newline(cleaned, settings.CHUNK_SIZE, settings.CHUNK_OVERLAP)

    docs = [d for d in docs_all if not _is_noisy(d.page_content)]

    for i, d in enumerate(docs):
        d.metadata = {"chunk_id": i}

    emb = get_embeddings()
    vs = FAISS.from_documents(docs, embedding=emb)
    (base / "faiss_index").mkdir(exist_ok=True)
    vs.save_local(str(base / "faiss_index"))

    (base / "chunks.count").write_text(str(len(docs)), encoding="utf-8")
    if docs:
        (base / "first.chunk.txt").write_text(docs[0].page_content, encoding="utf-8")
    return str(base)


# ========================= Поиск и ответ =========================

def _load_faiss(file_dir: str):
    emb = get_embeddings()
    vs = FAISS.load_local(str(Path(file_dir) / "faiss_index"), embeddings=emb, allow_dangerous_deserialization=True)
    return vs, emb

QA_PROMPT = (
    "Ты — помощник по русскоязычным договорам. Отвечай ТОЛЬКО по данным фрагментам. "
    "Ищи ответ внимательно, обращай внимание на слова из вопроса. " 
    "Если ответа нет — напиши: \"В документе не найдено\". Отвечай кратко.\n\n"
    "Вопрос: {question}\n\n"
    "Фрагменты:\n{context}\n"
)

def answer_question(file_dir: str, question: str, llm):
    vs, _ = _load_faiss(file_dir)
    question = _expand_query(question)
    try:
        docs = vs.max_marginal_relevance_search(
            question, k=settings.RETRIEVAL_K, fetch_k=max(settings.RETRIEVAL_K * 3, 24)
        )
    except Exception:
        docs = vs.similarity_search(question, k=settings.RETRIEVAL_K)

    parts: List[str] = []
    total = 0
    for d in docs:
        txt = (d.page_content or "").strip()
        txt = _squash_newlines_to_space(txt)
        if not txt:
            continue
        if total + len(txt) > settings.CONTEXT_CHARS_LIMIT:
            remain = settings.CONTEXT_CHARS_LIMIT - total
            if remain > 200:
                parts.append(txt[:remain])
            break
        parts.append(txt)
        total += len(txt)

    context = "\n \n".join(parts) if parts else "(пусто)"
    references = [{"rank": i + 1, "snippet": _squash_newlines_to_space((d.page_content or "")[:400])} for i, d in enumerate(docs)]

    prompt = QA_PROMPT.format(question=question, context=context)
    answer = llm.generate(prompt)
    return answer, references
